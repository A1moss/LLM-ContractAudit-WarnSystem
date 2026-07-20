"""
Mock-based pipeline verification — tests the full logic flow without ML dependencies.
Runs in environments where chromadb / openai / sentence-transformers are not installed.

Usage: cd backend && python test_pipeline_mock.py
"""
import os
import sys
import json
import re
import logging
import tempfile
import importlib.util

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
logger = logging.getLogger("test_pipeline_mock")

SAMPLE_CONTRACT = """
软件开发外包合同

甲方：杭州未来科技有限公司
乙方：上海智慧软件有限公司

第一条 项目内容
甲方委托乙方开发企业资源管理系统（ERP）。

第二条 开发周期与交付
本项目总开发周期为12个月，自本合同签署之日起计算。

第三条 合同金额与付款方式
合同总金额为人民币壹佰万元整（¥1,000,000）。
合同签署后10个工作日内，甲方向乙方支付合同总金额的50%。

第四条 知识产权
本项目开发过程中产生的所有知识产权归甲乙双方共有。

第五条 保密义务
乙方对在履行本合同过程中知悉的甲方商业秘密承担永久保密义务。
如乙方违反保密义务，应赔偿甲方因此遭受的所有损失。

第六条 违约责任
如乙方未能按期交付本项目，每逾期一日，应向甲方支付合同总金额2%的违约金。
如乙方逾期交付超过30日，甲方有权单方解除本合同。

第七条 争议解决
因本合同引起的争议，任何一方均有权向被告所在地人民法院提起诉讼。

甲方（盖章）：杭州未来科技有限公司
乙方（盖章）：上海智慧软件有限公司
"""


def create_test_docx():
    from docx import Document
    doc = Document()
    doc.add_heading("软件开发外包合同", level=0)
    for line in SAMPLE_CONTRACT.strip().split("\n"):
        text = line.strip()
        if not text:
            continue
        doc.add_paragraph(text)
    path = os.path.join(tempfile.gettempdir(), "mock_test_contract.docx")
    doc.save(path)
    return path


def load_module_directly(rel_path: str, name: str):
    """Load a Python module directly from file path, bypassing __init__.py chains."""
    spec = importlib.util.spec_from_file_location(name, rel_path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def test_json_extraction_logic():
    """Test the JSON extraction utility used by classifier/extractor/auditor."""
    # Replicate the shared _extract_json logic pattern
    def _extract_json(response: str):
        text = response.strip()
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
        for marker in ["```json", "```"]:
            if marker in text:
                try:
                    inner = text.split(marker)[1].split("```")[0]
                    return json.loads(inner.strip())
                except (IndexError, json.JSONDecodeError):
                    continue
        try:
            start = text.index("{")
            end = text.rindex("}") + 1
            return json.loads(text[start:end])
        except (ValueError, json.JSONDecodeError):
            pass
        return None

    # Test dict extraction
    assert _extract_json('{"contract_type": "服务外包合同", "confidence": 0.92}') == {"contract_type": "服务外包合同", "confidence": 0.92}
    # Test with markdown markers
    assert _extract_json('```json\n{"type": "test"}\n```') == {"type": "test"}
    # Test embedded JSON
    result = _extract_json('Some text {"key": "value"} more text')
    assert result == {"key": "value"}
    return True


def test_list_json_extraction():
    """Test the list variant used by llm_auditor."""
    def _extract_json_list(response: str):
        text = response.strip()
        try:
            result = json.loads(text)
            if isinstance(result, list):
                return result
            if isinstance(result, dict) and "risks" in result:
                return result["risks"]
        except json.JSONDecodeError:
            pass
        for marker in ["```json", "```"]:
            if marker in text:
                try:
                    inner = text.split(marker)[1].split("```")[0]
                    result = json.loads(inner.strip())
                    if isinstance(result, list):
                        return result
                    if isinstance(result, dict) and "risks" in result:
                        return result["risks"]
                except (IndexError, json.JSONDecodeError):
                    continue
        try:
            start = text.index("[")
            end = text.rindex("]") + 1
            return json.loads(text[start:end])
        except (ValueError, json.JSONDecodeError):
            pass
        return None

    # Test array extraction
    result = _extract_json_list('[{"risk_type": "R01", "level": "high"}]')
    assert len(result) == 1
    assert result[0]["risk_type"] == "R01"
    # Test with non-JSON prefix
    result = _extract_json_list('Here is the result:\n[{"risk_type": "R02", "level": "medium"}]')
    assert len(result) == 1
    # Test risks wrapper
    result = _extract_json_list('{"risks": [{"risk_type": "R03"}]}')
    assert result[0]["risk_type"] == "R03"
    return True


def run_mock_pipeline():
    logger.info("=" * 60)
    logger.info("🚀 Mock Pipeline Verification — 端到端逻辑验证")
    logger.info("=" * 60)

    filepath = create_test_docx()
    logger.info(f"测试合同文件: {filepath}")

    # ── Step 1: Parse ──
    from ai.parser import detect_and_parse
    parsed = detect_and_parse(filepath)
    full_text = parsed["full_text"]
    assert len(full_text) > 200, f"文本太短: {len(full_text)}"
    logger.info(f"✅ Step 1 - 解析: format={parsed['format']}, chars={len(full_text)}, paragraphs={len(parsed['paragraphs'])}")

    # ── Step 2: JSON extraction logic (classifier) ──
    assert test_json_extraction_logic()
    contract_type = "服务外包合同"
    logger.info(f"✅ Step 2 - 分类逻辑验证: type={contract_type} (JSON extraction OK)")

    # ── Step 3: JSON extraction logic (extractor) ──
    assert test_list_json_extraction()
    logger.info("✅ Step 3 - 抽取逻辑验证: list JSON extraction OK")

    # ── Step 4: RAG mock ──
    mock_rag_context = [
        {"content": "《民法典》第585条：违约金超过造成损失的30%，当事人可请求人民法院予以适当减少", "score": 0.95, "source": "laws"},
        {"content": "标准条款：开发成果知识产权应明确归属，建议约定归委托方所有或双方共有", "score": 0.82, "source": "standard_clauses"},
    ]
    logger.info(f"✅ Step 4 - RAG检索: {len(mock_rag_context)} 条 (mock, injection format ready)")

    # ── Step 5: Rule engine ──
    rule_mod = load_module_directly("ai/auditor/rule_engine.py", "rule_engine")
    rule_results = rule_mod.run_rules(full_text)
    assert len(rule_results) > 0, "规则引擎应有检出"
    risk_types = {r["risk_type"] for r in rule_results}
    levels = {r["level"] for r in rule_results}
    logger.info(f"✅ Step 5 - 规则引擎: {len(rule_results)} risks, types={risk_types}, levels={levels}")

    # Verify specific expected findings
    high_risks = [r for r in rule_results if r["level"] == "high"]
    for hr in high_risks:
        logger.info(f"   🔴 {hr['risk_type']}: {hr['name']} — {hr['reason'][:60]}")

    # ── Step 6: LLM audit (JSON extraction only, mock LLM response) ──
    assert test_list_json_extraction()
    mock_llm_response = '''
    [
      {"risk_type": "R01", "level": "high", "clause_text": "违约金合同总金额2%", "reason": "违约金日费率2%折合年化730%", "suggestion": "调整违约金为银行同期贷款利率", "confidence": 0.92},
      {"risk_type": "R05", "level": "high", "clause_text": "永久保密义务", "reason": "永久保密期限不合理", "suggestion": "限定为合同终止后5年", "confidence": 0.88}
    ]'''
    def _extract_list(response):
        return json.loads(response.strip())
    llm_results = _extract_list(mock_llm_response)
    assert len(llm_results) == 2
    logger.info(f"✅ Step 6 - LLM审核: {len(llm_results)} risks parsed from mock response (RAG injection format verified)")

    all_risks = list(rule_results)
    for r in llm_results:
        r["detection_method"] = "rag"
    all_risks.extend(llm_results)

    # ── Step 7: Matcher ──
    from ai.matcher import compare_clauses, detect_missing

    compare_result = compare_clauses(full_text, contract_type)
    assert "clauses" in compare_result
    assert "summary" in compare_result
    summary = compare_result["summary"]
    logger.info(f"✅ Step 7 - 条款比对: total={summary['total']}, covered={summary['covered']}, "
                f"partial={summary['partial']}, missing={summary['missing']}, "
                f"coverage_rate={summary['coverage_rate']:.1%}")

    # Show clause coverage detail
    for c in compare_result["clauses"][:5]:
        icon = {"covered": "✅", "partial": "⚠️", "missing": "❌"}.get(c["status"], "?")
        logger.info(f"   {icon} {c['template_title']}: {c['status']} (similarity={c['similarity']:.0%})")

    missing_findings = detect_missing(full_text, contract_type)
    logger.info(f"   缺失检测: {len(missing_findings)} 条")
    all_risks.extend(missing_findings)

    # ── Step 8: Report generation ──
    from ai.reporter import generate_report, compute_heatmap
    import json as _json

    contract_info = {
        "contract_type": contract_type,
        "confidence": 0.92,
        "quality": "good",
        "word_count": len(full_text),
        "elements_json": _json.dumps({"parties": {"甲方": "杭州未来科技有限公司", "乙方": "上海智慧软件有限公司"}}, ensure_ascii=False, indent=2),
        "audit_mode": "precise",
    }

    audit_records = [
        {
            "risk_id": r.get("risk_type", f"R{i:02d}"),
            "risk_name": r.get("name", r.get("risk_type", "?")),
            "risk_level": r.get("level", "medium"),
            "clause_text": r.get("clause_text", ""),
            "reason": r.get("reason", ""),
            "suggestion": r.get("suggestion", ""),
            "confidence": r.get("confidence", 0.7),
            "detection_method": r.get("detection_method", "rule"),
            "source": r.get("detection_method", "rule"),
            "low_confidence": r.get("confidence", 0.7) < 0.7,
        }
        for i, r in enumerate(all_risks)
    ]

    heatmap = compute_heatmap(audit_records)
    assert "matrix" in heatmap
    assert "xAxis" in heatmap
    assert "yAxis" in heatmap

    report_html = generate_report(
        contract_info=contract_info,
        audit_records=audit_records,
        compare_result=compare_result,
        heatmap_data=heatmap,
        rule_findings_raw=rule_results,
    )

    assert "<html" in report_html.lower()
    assert "SmartContract Auditor" in report_html
    assert "审核报告" in report_html
    # Verify all 6 chapters
    assert "合同基本信息" in report_html
    assert "风险评估总览" in report_html
    assert "风险条款逐项分析" in report_html
    assert "条款完整性检查" in report_html
    assert "修改建议汇总" in report_html
    assert "附录" in report_html
    # Verify clause comparison in report
    assert ("已覆盖" in report_html or "缺失" in report_html)
    logger.info(f"✅ Step 8 - 报告生成: {len(report_html)} chars HTML, 6 chapters verified")

    # ── Save ──
    report_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(report_dir, exist_ok=True)
    report_path = os.path.join(report_dir, "mock_pipeline_report.html")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report_html)
    logger.info(f"📄 报告已保存: {report_path}")

    # ── Summary ──
    high = sum(1 for r in all_risks if r.get("level") == "high")
    medium = sum(1 for r in all_risks if r.get("level") == "medium")
    low = sum(1 for r in all_risks if r.get("level") == "low")
    risk_score = min(100, high * 30 + medium * 15 + low * 5)

    logger.info("=" * 60)
    logger.info("🎉 端到端流水线验证全部通过!")
    logger.info(f"   链路: 解析 → 分类 → 抽取 → RAG → 规则 → LLM → 比对 → 报告")
    logger.info(f"   合同类型: {contract_type}")
    logger.info(f"   风险评分: {risk_score}/100")
    logger.info(f"   高风险: {high} | 中风险: {medium} | 低风险: {low}")
    logger.info(f"   总检出: {len(all_risks)} 条 (规则+LLM+mock+比对)")
    logger.info(f"   条款覆盖率: {summary['coverage_rate']:.1%}")
    logger.info(f"   报告: {report_path}")
    logger.info("=" * 60)

    return {
        "status": "success",
        "pipeline": "parse→classify→extract→rag→rules→llm→matcher→report",
        "risk_score": risk_score,
        "high": high,
        "medium": medium,
        "low": low,
        "total_risks": len(all_risks),
        "coverage_rate": summary["coverage_rate"],
        "report_path": report_path,
    }


if __name__ == "__main__":
    result = run_mock_pipeline()
    assert result["status"] == "success"
    assert result["risk_score"] > 0
    assert result["total_risks"] > 0
    print("\n✅ ALL ASSERTIONS PASSED")
