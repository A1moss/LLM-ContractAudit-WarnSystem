"""
端到端流水线测试脚本
运行方式: cd backend && python test_pipeline.py

测试链路: 合同上传 → 解析 → 分类 → 要素抽取 → RAG检索 → 规则审核 → LLM审核 → 条款比对 → 报告生成
"""
import os
import sys
import json
import logging
import tempfile

# 确保 backend 在 path 中
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
logger = logging.getLogger("test_pipeline")

# ── 样本合同文本 ──
SAMPLE_CONTRACT = """
软件开发外包合同

甲方：杭州未来科技有限公司（以下简称"甲方"）
地址：浙江省杭州市西湖区文三路478号华星时代广场
法定代表人：张三

乙方：上海智慧软件有限公司（以下简称"乙方"）
地址：上海市浦东新区张江高科技园区碧波路690号
法定代表人：李四

鉴于甲方需要开发一套企业资源管理系统（ERP），乙方具备相应的技术开发能力，经双方友好协商，达成如下协议：

第一条 项目内容
1.1 甲方委托乙方开发企业资源管理系统（以下简称"本项目"），包括以下模块：
（1）采购管理模块
（2）库存管理模块
（3）财务管理模块
（4）人力资源管理模块
1.2 详细功能需求见附件一《功能需求规格说明书》。

第二条 开发周期与交付
2.1 本项目总开发周期为12个月，自本合同签署之日起计算。
2.2 乙方应在2025年3月15日前完成全部开发工作并交付验收。
2.3 交付内容包括：全部源代码、技术文档、用户手册、部署脚本。

第三条 合同金额与付款方式
3.1 本项目合同总金额为人民币壹佰万元整（¥1,000,000），含税。
3.2 付款方式如下：
（1）合同签署后10个工作日内，甲方向乙方支付合同总金额的50%，即人民币伍拾万元整（¥500,000）；
（2）项目中期验收通过后10个工作日内，甲方向乙方支付合同总金额的30%，即人民币叁拾万元整（¥300,000）；
（3）项目最终验收通过后10个工作日内，甲方向乙方支付合同总金额的20%，即人民币贰拾万元整（¥200,000）。

第四条 知识产权
4.1 本项目开发过程中产生的所有知识产权（包括但不限于源代码、文档、设计图纸）归甲乙双方共有。
4.2 未经甲方书面同意，乙方不得将本项目相关技术用于其他项目。

第五条 保密义务
5.1 乙方对在履行本合同过程中知悉的甲方商业秘密、技术资料和经营信息承担永久保密义务。
5.2 保密义务不因本合同的终止或解除而终止。
5.3 如乙方违反保密义务，应赔偿甲方因此遭受的所有损失。

第六条 违约责任
6.1 如乙方未能按期交付本项目，每逾期一日，应向甲方支付合同总金额2%的违约金。
6.2 如乙方交付的项目不符合验收标准，应无条件进行修改直至通过验收。
6.3 如乙方逾期交付超过30日，甲方有权单方解除本合同，并要求乙方退还已支付的全部款项。

第七条 争议解决
7.1 因本合同引起的争议，双方应协商解决；协商不成的，任何一方均有权向被告所在地人民法院提起诉讼。

第八条 合同变更与解除
8.1 任何一方不得擅自变更或解除本合同。
8.2 如需变更合同内容，须经双方协商一致并签署书面变更协议。

第九条 签署与生效
9.1 本合同自双方签字盖章之日起生效。
9.2 本合同一式两份，甲乙双方各执一份，具有同等法律效力。

甲方（盖章）：杭州未来科技有限公司
授权代表（签字）：________
日期：2024年3月15日

乙方（盖章）：上海智慧软件有限公司
授权代表（签字）：________
日期：2024年3月15日
"""


def create_sample_contract() -> str:
    """创建样本合同 .docx 文件"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("软件开发外包合同", level=0)

        for paragraph_text in SAMPLE_CONTRACT.strip().split("\n"):
            line = paragraph_text.strip()
            if not line:
                continue
            if line.startswith("第") and "条" in line[:10]:
                doc.add_heading(line, level=1)
            elif line.startswith(("甲方", "乙方", "鉴于")):
                doc.add_paragraph(line, style="Normal")
            else:
                doc.add_paragraph(line, style="Normal")

        filepath = os.path.join(tempfile.gettempdir(), "test_contract_pipeline.docx")
        doc.save(filepath)
        logger.info(f"样本合同已创建: {filepath}")
        return filepath
    except ImportError:
        # Fallback: use a .txt file
        filepath = os.path.join(tempfile.gettempdir(), "test_contract_pipeline.txt")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(SAMPLE_CONTRACT)
        logger.warning("python-docx 不可用，使用 .txt 格式（需手动添加到 parser）")
        logger.info(f"样本合同已创建: {filepath}")
        return filepath


def run_pipeline_on_sample():
    """在样本合同上运行完整流水线"""
    from ai.pipeline import run_pipeline

    # 创建样本合同
    filepath = create_sample_contract()

    if not os.path.exists(filepath):
        logger.error(f"样本合同创建失败: {filepath}")
        return None

    logger.info("=" * 60)
    logger.info("🚀 启动端到端流水线测试")
    logger.info(f"   合同文件: {filepath}")
    logger.info(f"   审核模式: precise")
    logger.info("=" * 60)

    # 运行流水线
    result = run_pipeline(
        file_path=filepath,
        contract_type_hint="服务外包合同",
        audit_mode="precise",
    )

    # 输出结果摘要
    logger.info("=" * 60)
    logger.info("📊 流水线执行结果")
    logger.info(f"   状态: {result.get('status')}")
    logger.info(f"   合同类型: {result.get('contract_type')}")
    logger.info(f"   总耗时: {result.get('timing', {}).get('total_s', 0):.1f}s")

    # 各步骤耗时
    step_times = result.get("timing", {}).get("steps", {})
    logger.info("   各步骤耗时:")
    for step_name, duration in step_times.items():
        logger.info(f"     {step_name}: {duration:.1f}s")

    # 摘要
    summary = result.get("summary", {})
    logger.info(f"   风险评分: {summary.get('risk_score', 0)}/100")
    logger.info(f"   高风险: {summary.get('high', 0)}, 中风险: {summary.get('medium', 0)}, 低风险: {summary.get('low', 0)}")
    logger.info(f"   总风险数: {summary.get('total_risks', 0)}")
    logger.info(f"   条款覆盖率: {summary.get('coverage_rate', 0):.1%}")

    # 各步骤详情
    steps = result.get("steps", {})
    for step_name, step_info in steps.items():
        status = step_info.get("status", "?")
        icon = "✅" if status == "ok" else ("⚠️" if status == "warning" else ("❌" if status == "error" else "⏭️"))
        logger.info(f"   {icon} {step_name}: {status}")

    # 保存报告
    report_dir = os.path.join(os.path.dirname(__file__), "reports")
    os.makedirs(report_dir, exist_ok=True)
    report_path = os.path.join(report_dir, "pipeline_test_report.html")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(result.get("report_html", "<html><body>Report empty</body></html>"))
    logger.info(f"📄 审核报告已保存: {report_path}")

    # 保存完整结果 JSON
    json_path = os.path.join(report_dir, "pipeline_test_result.json")
    result_for_json = {
        "status": result.get("status"),
        "contract_type": result.get("contract_type"),
        "audit_mode": result.get("audit_mode"),
        "summary": summary,
        "timing": result.get("timing"),
        "steps": {k: v for k, v in steps.items()},
        "all_risks_count": len(result.get("all_risks", [])),
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(result_for_json, f, ensure_ascii=False, indent=2)
    logger.info(f"📋 完整结果已保存: {json_path}")

    return result


if __name__ == "__main__":
    run_pipeline_on_sample()
